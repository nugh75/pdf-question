# %%
# !pip install langchain
# !pip install openai
#!pip install PyPDF2
# !pip install faiss-cpu
# !pip install tiktoken
#!pip install --upgrade pip
!pip install -U langchain-community
#!pip install python-dotenv

# %%

!pip freeze > requirements.txt


# %%
from PyPDF2 import PdfReader
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import FAISS

# %%
from dotenv import load_dotenv
import os

# Carica le variabili di ambiente dal file .env
load_dotenv()

# Ora puoi accedere alle variabili di ambiente
openai_api_key = os.getenv("OPENAI_API_KEY")

# %%
# provide the path of  pdf file/files.
pdfreader = PdfReader('hoder.pdf')

# %%
from typing_extensions import Concatenate
# read text from pdf
raw_text = ''
for i, page in enumerate(pdfreader.pages):
    content = page.extract_text()
    if content:
        raw_text += content

# %%
raw_text

# %%
# We need to split the text using Character Text Split such that it sshould not increse token size
text_splitter = CharacterTextSplitter(
    separator = "\n",
    chunk_size = 800,
    chunk_overlap  = 200,
    length_function = len,
)
texts = text_splitter.split_text(raw_text)

# %%
len(texts)

# %%
# Download embeddings from OpenAI
embeddings = OpenAIEmbeddings()

# %%
document_search = FAISS.from_texts(texts, embeddings)

# %%
document_search


# %%
from langchain.chains.question_answering import load_qa_chain
from langchain.llms import OpenAI

# %%
chain = load_qa_chain(OpenAI(), chain_type="stuff")

# %%
query = "Quali sono gli obiettivi della ricerca? scrivi almeno tre capoversi"
docs = document_search.similarity_search(query)
output = chain.run(input_documents=docs, question=query)  # Salva l'output in una variabile

# Scrivi la domanda e la risposta nel file, separati da una riga vuota, e metti la domanda tra ##
with open('output.txt', 'w', encoding='utf-8') as file:
    file.write('## ' + query + ' ##' + '\n' + output)

# %%
# Seconda domanda
query2 = "Quali sono le metodologie della ricerca? alemeno tre capoversi"
docs2 = document_search.similarity_search(query2)
output2 = chain.run(input_documents=docs2, question=query2)

# Aggiungi la seconda domanda e la risposta al file, separati da una riga vuota, e metti la domanda tra ##
with open('output.txt', 'a', encoding='utf-8') as file:
    file.write('## ' + query2 + ' ##' + '\n' + output2 + '\n')

# %%
#terza domanda
query2 = "Quali sono i risultati della ricerca? alemeno tre capoversi"
docs2 = document_search.similarity_search(query2)
output2 = chain.run(input_documents=docs2, question=query2)

# Aggiungi la seconda domanda e la risposta al file, separati da una riga vuota, e metti la domanda tra ##
with open('output.txt', 'a', encoding='utf-8') as file:
    file.write('## ' + query2 + ' ##' + '\n' + output2 + '\n')

# %%
from docx import Document

# Crea un nuovo documento Word
doc = Document()

# Leggi il contenuto del file output.txt
with open('output.txt', 'r', encoding='utf-8') as file:
    content = file.read()

# Dividi il contenuto in blocchi usando le righe vuote come separatori
blocks = content.split('\n\n')

# Itera sui blocchi e aggiungi al documento Word
for i in range(0, len(blocks), 2):
    if i < len(blocks):
        question = blocks[i].strip()
        if question.startswith('##') and question.endswith('##'):
            question = question[2:-2].strip()  # Rimuovi i ## dalla domanda
            doc.add_heading(question, level=1)  # Aggiungi la domanda come titolo

    if i + 1 < len(blocks):
        answer = blocks[i + 1].strip()
        doc.add_paragraph(answer)  # Aggiungi la risposta come paragrafo

# Salva il documento Word
doc.save('output.docx')


